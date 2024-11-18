import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor, Pt
from typing import Optional, Dict, Union, Any
from pathlib import Path

class TableDocxHandler:
    def __init__(self, doc):
        self.doc = Document(doc)

    def __repr__(self):
        return f"TableDocxHandler(doc with {len(self.doc.tables)} tables)"

    def __len__(self):
        return len(self.doc.tables)

    def __getitem__(self, index):
        return self.doc.tables[index]

    def __setitem__(self, index, df: pd.DataFrame):
        # replace table in document
        table = self.doc.tables[index]
        self.clear_table(table)
        self.insert_dataframe(table, df)

    def clear_table(self, table):
        for row in table.rows:
            table._element.remove(row._element)

    def insert_dataframe(
            self, table, df: pd.DataFrame, style: Optional[Dict] = None
    ) -> None:
        for i, col in enumerate(df.columns):
            cell = table.cell(0, i)
            cell.text = col
            if style:
                self.apply_cell_style(cell, style, is_header=True)

        for _, row in df.iterrows():
            cells = table.add_row().cells
            for i, val in enumerate(row):
                cells[i].text = str(val)
                if style:
                    self.apply_cell_style(cells[i], style, is_header=False)

    def add_table(
        self, df: pd.DataFrame, title: Optional[str] = None, style: Optional[Dict] = None
    ) -> None:
        if title:
            p = self.doc.add_paragraph(
            title, style='Normal'
            )
            p.alignment=WD_ALIGN_PARAGRAPH.LEFT
            p.runs[0].font.bold = True

        table = self.doc.add_table(rows=1, cols=len(df.columns))
        self.set_table_borders(table)
        self.insert_dataframe(table, df)

    def set_table_borders(self, table):
        tbl = table._tbl
        tbl_pr = tbl.tblPr
        tbl_borders = OxmlElement('w:tblBorders')

        for border_type in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_type}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4') #border size
            tbl_borders.append(border)

        tbl_pr.append(tbl_borders)

    def set_cell_margins(self, cell, top=0, bottom=0, start=0, end=0):
        tc = cell._element
        tcPr = tc.get_or_add_tcPr()
        tcMar = OxmlElement('w:tcMar')

        for type, val in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
            node = OxmlElement(f'w:{type}')
            node.set(qn('w:w'), str(val))
            node.set(qn('w:type'), 'dxa')
            tcMar.append(node)

        tcPr.append(tcMar)

    def apply_style(self, run, style):
        # uses a style dictionary to apply styles to a run
        run.font.bold = style.get('bold', False)
        run.font.size = style.get('size', 11)
        run.font.name = style.get('name', 'Calibri')
        if 'color' in style:
            run.font.color.rgb = RGBColor.from_string(style['color'])

    def apply_cell_style(self, cell, style, is_header=False):
        # Apply styles to table cells, optionally differentiating headers
        p = cell.paragraphs[0]
        run = p.runs[0] if p.runs else p.add_run()
        if is_header:
            run.font.bold = style.get('header_bold', True)
        else:
            run.font.bold = style.get('bold', False)
        run.font.size = Pt(style.get('size', 11))
        run.font.name = style.get('font', 'Calibri')
        p.alignment = getattr(WD_ALIGN_PARAGRAPH, style.get('alignment', 'LEFT').upper())
