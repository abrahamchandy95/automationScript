import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Inches
from typing import Optional, Dict, List, Union
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os
from pathlib import Path

from .tableDocHandler import TableDocxHandler

class DocxTemplator:
    def __init__(self) -> None:

        self.doc = Document()
        self.table_handler = TableDocxHandler(self.doc)
        self.save_path = None
        self.content = []

    def __repr__(self)-> str:
        return f"DocxGenerator with {len(self.doc.tables)}"

    def __len__(self)-> int:
        return len(self.doc.tables)

    def __getitem__(self, index: int):
        return self.doc.tables[index]

    def __setitem__(self, index: int, df: pd.DataFrame) -> None:
        # allows to replace a table at index with a new DataFrame
        self.table_handler[index] = df

    def __enter__(self):
        # allows use in a with statement
        return self

    def set_save_path(self, path):
        self.save_path = path

    def add_content(
            self,
            content_type: str,
            content: Union[str, Path, pd.DataFrame],
            style: Dict [str, str] = {}
    ):
        self.content.append(
            (content_type, content, style if style else {})
        )

    def add_norm_para(self, content: str, style: Dict[str, str])-> None:
        p = self.doc.add_paragraph()
        run = p.add_run(content)
        run.font.bold = bool(style.get('bold', False))
        run.font.size = Pt(int(style.get('size', 11)))
        run.font.name = style.get('font', 'Calibri')
        if 'color' in style:
            run.font.color.rgb = RGBColor.from_string(style['color'])
        if style.get('underline', False):
            run.font.underline = True
        p.alignment = getattr(
            WD_ALIGN_PARAGRAPH, style.get('alignment', 'LEFT').upper()
        )
        space_after = Pt(int(style.get('space_after', 4)))
        p.paragraph_format.space_after = space_after

    def add_bold_prefix_para(
            self, prefix: str, content: str, style: Dict[str, str]
    )-> None:
        p = self.doc.add_paragraph()
        run = p.add_run(prefix)
        run.font.bold = True
        run.font.size = Pt(int(style.get('font_size', 11)))
        run.font.name = style.get('font', 'Calibri')
        run.font.color.rgb = RGBColor.from_string(style.get('color', '000000'))

        run =p.add_run(content)
        run.bold = False
        run.font.size = Pt(int(style.get('font_size', 11)))
        run.font.name = 'Calibri'
        run.font.color.rgb = RGBColor.from_string(style.get('color', '000000'))

    def add_image(self, image_path_str: str, title: Optional[str]=None):
        image_path = Path(image_path_str)
        if title:
            self.doc.add_paragraph(title, style='Normal')

        exts = ['.jpg', '.jpeg', '.png', '.gif']
        img_path = None
        img_path = next(
            (image_path.with_suffix(e) for e in exts \
                if image_path.with_suffix(e).exists()), None
        )
        if img_path:
            self.doc.add_picture(str(img_path), width=Inches(6))
            return

    def generate_report(self)-> None:

        for content_type, data, style in self.content:
            if content_type == 'paragraph':
                self.add_norm_para(data, style)

            elif content_type == 'bold_prefix_paragraph':
                self.add_bold_prefix_para(data['prefix'], data['content'], style)

            elif content_type == 'image':
                self.add_image(data)

            elif content_type == 'table':
                self.table_handler.add_table(data, style=style)

            elif content_type == 'page_break':
                self.doc.add_page_break()

    def try_add_image(self, image_path):
        extensions = ['.jpg', '.jpeg', '.png']
        image_file = Path(image_path)
        if image_file.exists():
            self.doc.add_picture(str(image_file), width=Inches(6))
            self.doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            return True
        else:
            for ext in extensions:
                mod_img_path = image_file.with_suffix(ext)
                if mod_img_path.exists():
                    self.doc.add_picture(str(mod_img_path), width=Inches(6))
                    self.doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    return True
        return False


    def save_doc(self)-> None:
        if not self.save_path:
            raise ValueError("Save path not set")
        self.doc.save(self.save_path)
